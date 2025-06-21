import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


class BillDocumentGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("Bill Document Generator")
        self.root.geometry("1000x700")

        # Configuration variables (easily adjustable)
        self.document_orientation = "landscape"  # or "portrait"
        self.screenshot_width = 4.0  # in inches (larger since only 2 per page)
        self.screenshot_height = 6.0  # in inches
        self.horizontal_spacing = 1.0  # in inches between screenshots
        self.vertical_spacing = 1.0  # in inches between pages
        self.page_margins = 0.5  # in inches
        self.screenshots_per_page = 2  # Only 2 screenshots per page
        self.font_size = 12  # in points for header text

        # Create variables to hold image paths
        self.image_paths = [None] * 8
        self.image_previews = [None] * 8

        self.setup_ui()

    def setup_ui(self):
        # Main frame
        main_frame = tk.Frame(self.root, padx=20, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Title
        title_label = tk.Label(main_frame, text="Bill Document Generator", font=("Arial", 16, "bold"))
        title_label.grid(row=0, column=0, columnspan=4, pady=(0, 20))

        # Instructions
        instructions = tk.Label(main_frame,
                                text="Upload screenshots of bills for each room (up to 8).\nDocument will have 2 screenshots per page.",
                                wraplength=600)
        instructions.grid(row=1, column=0, columnspan=4, pady=(0, 20))

        # Room frames for image upload
        self.room_frames = []
        for i in range(8):
            room_frame = tk.Frame(main_frame, bd=2, relief=tk.GROOVE, padx=10, pady=10)
            room_frame.grid(row=2 + (i // 4), column=i % 4, padx=5, pady=5, sticky="nsew")

            room_label = tk.Label(room_frame, text=f"Room {i + 1}", font=("Arial", 10, "bold"))
            room_label.pack()

            preview_label = tk.Label(room_frame, text="No image selected", width=20, height=5,
                                     relief=tk.SUNKEN, bg="white")
            preview_label.pack()

            btn = tk.Button(room_frame, text="Upload Screenshot",
                            command=lambda idx=i: self.upload_image(idx))
            btn.pack(pady=(5, 0))

            self.room_frames.append((room_frame, preview_label))

            # Configure grid weights for even spacing
            main_frame.columnconfigure(i % 4, weight=1)

        # Generate button
        generate_btn = tk.Button(main_frame, text="Generate Document",
                                 command=self.generate_document, bg="#4CAF50", fg="white",
                                 font=("Arial", 12, "bold"))
        generate_btn.grid(row=4, column=0, columnspan=4, pady=20)

        # Progress bar
        self.progress = ttk.Progressbar(main_frame, orient=tk.HORIZONTAL, length=300, mode='determinate')
        self.progress.grid(row=5, column=0, columnspan=4, pady=(0, 10))

        # Status label
        self.status_label = tk.Label(main_frame, text="", fg="blue")
        self.status_label.grid(row=6, column=0, columnspan=4)

    def upload_image(self, index):
        file_path = filedialog.askopenfilename(
            title=f"Select bill screenshot for Room {index + 1}",
            filetypes=[("Image files", "*.png *.jpg *.jpeg"), ("All files", "*.*")]
        )

        if file_path:
            try:
                # Store the image path
                self.image_paths[index] = file_path

                # Display a thumbnail preview
                img = Image.open(file_path)
                img.thumbnail((150, 150))
                photo = ImageTk.PhotoImage(img)

                # Update the preview label
                preview_label = self.room_frames[index][1]
                preview_label.config(image=photo, text="")
                preview_label.image = photo  # Keep a reference

            except Exception as e:
                messagebox.showerror("Error", f"Could not load image: {str(e)}")

    def generate_document(self):
        # Check if at least one image is selected
        if not any(self.image_paths):
            messagebox.showwarning("Warning", "Please upload at least one bill screenshot.")
            return

        # Ask for save location
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
            title="Save Document As"
        )

        if not save_path:
            return  # User cancelled

        self.progress['value'] = 10
        self.root.update_idletasks()

        try:
            # Create a new Word document
            doc = Document()

            # Set page orientation
            if self.document_orientation == "landscape":
                section = doc.sections[0]
                section.orientation = WD_ORIENT.LANDSCAPE
                new_width, new_height = section.page_height, section.page_width
                section.page_width = new_width
                section.page_height = new_height

            # Set margins
            for section in doc.sections:
                section.top_margin = Inches(self.page_margins)
                section.bottom_margin = Inches(self.page_margins)
                section.left_margin = Inches(self.page_margins)
                section.right_margin = Inches(self.page_margins)

            # Add a title
            doc.add_heading('Monthly Bills Summary', level=1)

            self.progress['value'] = 30
            self.root.update_idletasks()

            # Process each image and add to document
            images_added = 0
            current_page_images = 0

            for i, img_path in enumerate(self.image_paths):
                if img_path:
                    # Start a new page if needed
                    if current_page_images % self.screenshots_per_page == 0:
                        if images_added > 0:  # Don't add page break for first image
                            doc.add_page_break()

                        # Add page header
                        para = doc.add_paragraph()
                        run = para.add_run(f"Bills (Page {images_added // self.screenshots_per_page + 1})")
                        run.bold = True
                        run.font.size = Pt(14)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()

                    try:
                        # Create a table to hold the two images side by side
                        if current_page_images % self.screenshots_per_page == 0:
                            table = doc.add_table(rows=1, cols=2)
                            table.autofit = False

                            # Set column widths
                            for cell in table.columns[0].cells:
                                cell.width = Inches(self.screenshot_width)
                            for cell in table.columns[1].cells:
                                cell.width = Inches(self.screenshot_width)

                            # Add spacing between columns
                            table.columns[0].cells[0].width = Inches(self.screenshot_width)
                            table.columns[1].cells[0].width = Inches(self.screenshot_width)

                        # Add image to the table
                        cell = table.rows[0].cells[current_page_images % self.screenshots_per_page]
                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run()

                        # Add room number label above the image
                        room_label = paragraph.add_run(f"Room {i + 1}\n")
                        room_label.bold = True
                        room_label.font.size = Pt(self.font_size)

                        # Add the image
                        run.add_picture(img_path,
                                        width=Inches(self.screenshot_width - 0.5),
                                        height=Inches(self.screenshot_height - 0.5))

                        # Center the image in the cell
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        images_added += 1
                        current_page_images += 1

                    except Exception as e:
                        messagebox.showwarning("Warning", f"Could not add image for Room {i + 1}: {str(e)}")

                    self.progress['value'] = 30 + (i * 70 / 8)
                    self.root.update_idletasks()

            # Save the document
            doc.save(save_path)

            self.progress['value'] = 100
            self.status_label.config(text=f"Document successfully saved to: {save_path}", fg="green")

            messagebox.showinfo("Success", "Document generated successfully!")

        except Exception as e:
            self.progress['value'] = 0
            self.status_label.config(text=f"Error: {str(e)}", fg="red")
            messagebox.showerror("Error", f"Could not generate document: {str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = BillDocumentGenerator(root)
    root.mainloop()